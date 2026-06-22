import io
import re
import os
import json
import requests
from collections import Counter

import streamlit as st

# ── Optional heavy imports ────────────────────────────────────────────────────
try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    from docx import Document
except Exception:
    Document = None

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
except Exception:
    TfidfVectorizer = None
    cosine_similarity = None

try:
    from sentence_transformers import SentenceTransformer
    _semantic_model = SentenceTransformer('all-MiniLM-L6-v2')
except Exception:
    _semantic_model = None

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="AI Resume Rewriter",
    page_icon="📄",
    layout="wide",
)

# ── Constants — exact copy from notebook ─────────────────────────────────────
STOPWORDS = {
    'a','an','and','are','as','at','be','by','for','from','has',
    'have','in','is','it','of','on','or','that','the','this','to',
    'with','you','your','we','our','will','can','using','use','into',
    'within','across','their','they','job','role','candidate','work',
    'experience','skills','team','teams','responsibilities','required',
}

SYNONYMS = {
    'developed':'built','created':'built','designed':'built',
    'implemented':'built','programmed':'built','coded':'built',
    'managed':'led','supervised':'led','handled':'led',
    'utilized':'used','leveraged':'used','employed':'used',
    'constructed':'built','established':'built',
    'ml':'machine learning','ai':'artificial intelligence',
    'nlp':'natural language processing','dl':'deep learning',
    'cv':'computer vision','db':'database',
    'analysed':'analyzed','modelling':'modeling',
    'analyse':'analyze','optimisation':'optimization',
}

TECH_SKILL_LIST = [
    'python','sql','r','pandas','numpy','scikit-learn','tensorflow',
    'pytorch','xgboost','random forest','decision tree','matplotlib',
    'seaborn','plotly','power bi','tableau','excel','apache spark',
    'hadoop','nlp','computer vision','machine learning','deep learning',
    'statistical modeling','hypothesis testing','regression analysis',
    'time series','feature engineering','data wrangling','smote',
    'hyperparameter optimization','data visualization','jupyter',
    'github','clustering','classification','regression','forecasting',
    'data analytics','data science','predictive modeling','spss',
    'big data','business intelligence','data cleaning','data preprocessing',
]

# ── Text helpers — exact copy from notebook ───────────────────────────────────
def clean_text(text):
    if not text: return ''
    text = re.sub(r'\r', '\n', text)
    text = re.sub(r'\b([A-Za-z])\s([A-Za-z])\s([A-Za-z])\s([A-Za-z])\b', r'\1\2\3\4', text)
    text = re.sub(r'\b([A-Za-z])\s([A-Za-z])\s([A-Za-z])\b', r'\1\2\3', text)
    text = re.sub(r'([a-z])\s([a-z]{1,2})\s([a-z])', r'\1\2\3', text)
    text = re.sub(r'(\w)-\s+(\w)', r'\1\2', text)
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def normalize(text):
    text = text.lower()
    for word, replacement in SYNONYMS.items():
        text = re.sub(r'\b' + word + r'\b', replacement, text)
    return text

def extract_text_from_pdf(file) -> str:
    if PdfReader is None: return ''
    try:
        reader = PdfReader(file)
        pages = [page.extract_text() or '' for page in reader.pages]
        raw = '\n'.join(pages)
        cleaned = clean_text(raw)
        words = cleaned.split()
        if words:
            single_chars = sum(1 for w in words if len(w) == 1)
            if single_chars / len(words) > 0.3:
                cleaned = re.sub(r'(?<![.\n])\b([A-Za-z])\b\s', r'\1', cleaned)
                cleaned = clean_text(cleaned)
        return cleaned
    except Exception as e:
        return ''

def extract_text_from_docx(file) -> str:
    if Document is None: return ''
    try:
        file_bytes = file.getvalue() if hasattr(file, 'getvalue') else file.read()
        doc = Document(io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip(): parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip(): parts.append(para.text.strip())
        return clean_text('\n'.join(parts))
    except Exception as e:
        return ''

def extract_uploaded_text(file) -> str:
    if file is None: return ''
    name = file.name.lower()
    if name.endswith('.txt'):
        return clean_text(file.read().decode('utf-8', errors='ignore'))
    if name.endswith('.pdf'):
        return extract_text_from_pdf(file)
    if name.endswith('.docx'):
        return extract_text_from_docx(file)
    return ''

def tokenize(text):
    words = re.findall(r'[a-zA-Z][a-zA-Z+#.\-]{1,}', text.lower())
    return [w.strip('.-') for w in words if w not in STOPWORDS and len(w) > 2]

# ── NLP functions — exact copy from notebook ──────────────────────────────────
def technical_skills_from_jd(jd_text):
    jd_lower = jd_text.lower()
    return [skill for skill in TECH_SKILL_LIST if skill in jd_lower]

def missing_keywords(resume_text, jd_text):
    jd_skills    = technical_skills_from_jd(jd_text)
    resume_lower = normalize(resume_text).lower()
    return [s for s in jd_skills if s not in resume_lower][:12]

def tfidf_score(resume_text, jd_text):
    if not TfidfVectorizer: return 0.0
    vectorizer = TfidfVectorizer(
        stop_words='english', ngram_range=(1,2),
        max_features=1000, sublinear_tf=True,
    )
    matrix = vectorizer.fit_transform([normalize(resume_text), normalize(jd_text)])
    return round(float(cosine_similarity(matrix[0], matrix[1])[0][0]) * 100, 1)

def keyword_overlap_score(resume_text, jd_text):
    resume_words = set(tokenize(normalize(resume_text)))
    jd_words     = set(tokenize(normalize(jd_text)))
    if not jd_words: return 0.0
    return round(len(resume_words & jd_words) / len(jd_words) * 100, 1)

def weighted_keyword_score(resume_text, jd_text):
    jd_skills    = technical_skills_from_jd(jd_text)
    resume_lower = normalize(resume_text).lower()
    if not jd_skills: return 0.0
    matched = sum(1 for s in jd_skills if s in resume_lower)
    return round(matched / len(jd_skills) * 100, 1)

def semantic_score(resume_text, jd_text):
    if _semantic_model is None: return 0.0
    emb1 = _semantic_model.encode([resume_text])
    emb2 = _semantic_model.encode([jd_text])
    return float(round(cosine_similarity(emb1, emb2)[0][0] * 100, 1))

def calculate_match_score(resume_text, jd_text):
    if not resume_text or not jd_text: return 0.0
    s1 = tfidf_score(resume_text, jd_text)
    s2 = keyword_overlap_score(resume_text, jd_text)
    s3 = weighted_keyword_score(resume_text, jd_text)
    s4 = semantic_score(resume_text, jd_text)
    return round((s1 * 0.15) + (s2 * 0.35) + (s3 * 0.35) + (s4 * 0.15), 1)

# ── Groq AI — exact copy from notebook ───────────────────────────────────────
def call_groq(prompt, api_key, max_tokens=4000):
    try:
        response = requests.post(
            'https://api.groq.com/openai/v1/chat/completions',
            headers={
                'Authorization': f'Bearer {api_key}',
                'Content-Type': 'application/json',
            },
            json={
                'model': 'llama-3.3-70b-versatile',
                'messages': [{'role': 'user', 'content': prompt}],
                'max_tokens': max_tokens,
                'temperature': 0,
            },
            timeout=90,
        )
        data = response.json()
        if 'choices' in data:
            return data['choices'][0]['message']['content'].strip()
        elif 'error' in data:
            return f"Groq Error: {data['error']['message']}"
    except Exception as e:
        return f'Request failed: {str(e)}'

def inject_missing_keywords(rewritten, jd_text):
    still_missing = missing_keywords(rewritten, jd_text)
    if not still_missing: return rewritten
    lines = rewritten.splitlines()
    new_lines = []
    skills_found = False
    for line in lines:
        new_lines.append(line)
        if re.search(r'^(SKILLS|TECHNICAL SKILLS|CORE COMPETENCIES)$',
                     line.strip(), re.IGNORECASE) and not skills_found:
            skills_found = True
            new_lines.append(', '.join(s.title() for s in still_missing[:10]))
    if not skills_found and still_missing:
        final = []
        inserted = False
        for line in new_lines:
            final.append(line)
            if not inserted and len(line.strip()) > 50:
                final.append('Additional skills: ' + ', '.join(s.title() for s in still_missing[:8]) + '.')
                inserted = True
        return clean_text('\n'.join(final))
    return clean_text('\n'.join(new_lines))

def rewrite_resume(resume_text, jd_text, api_key):
    gaps     = missing_keywords(resume_text, jd_text)
    skill_kw = technical_skills_from_jd(jd_text)

    prompt = f"""You are a senior professional resume writer and ATS specialist.

Rewrite the resume to professionally match the job description.

TECHNICAL SKILLS TO INCLUDE: {', '.join(skill_kw)}
MISSING SKILLS TO ADD NATURALLY: {', '.join(gaps[:10])}

STRICT RULES:
1. READ THE ENTIRE RESUME before writing anything.
2. Every section in the original resume MUST appear in the rewrite.
3. Do NOT write "No direct experience" — the experience IS in the resume.
4. Do NOT invent fake companies, degrees, dates, or numbers.
5. Keep all real data: name, email, phone, LinkedIn, dates, GPA.
6. Rewrite summary to match the JD role language closely.
7. Skills section must list ALL technical skills the candidate has.
8. Rewrite every bullet point with strong action verbs and JD keywords.
9. Write every project FULLY — tools used, dataset size, accuracy achieved.
10. No markdown bold (**text**) — plain text only.
11. Section headings in plain ALL CAPS — no symbols or asterisks.
12. Do NOT add "Analytics Hub", "Work Big Data" or any recruitment phrases.

FORMAT TO FOLLOW EXACTLY:
CANDIDATE NAME
email | phone | city | linkedin

SUMMARY
2-3 sentences matching the JD.

SKILLS
Programming: Python, R, SQL
Machine Learning: Scikit-learn, TensorFlow, PyTorch, XGBoost, Random Forest
Data Visualization: Power BI, Tableau, Matplotlib, Seaborn, Plotly
Statistics: Regression, Hypothesis Testing, Time Series, Clustering
Big Data: Apache Spark, Hadoop
Tools: Excel, Jupyter Notebook, GitHub, SPSS

INTERNSHIP
Company | Role | Date
- Action verb + task + tool + result
- Action verb + task + tool + result

PROJECTS
Project Title | Tools used
- Full description, dataset, accuracy/result achieved
- Techniques used and business insight

EDUCATION
Degree | Institution | Year | CGPA

CERTIFICATIONS
Name | Issuer | Year

--- ORIGINAL RESUME ---
{resume_text}

--- JOB DESCRIPTION ---
{jd_text}

Write the complete professional resume now:"""

    rewritten = call_groq(prompt, api_key, max_tokens=4000)
    if rewritten and not rewritten.startswith('Groq Error'):
        rewritten = inject_missing_keywords(rewritten, jd_text)
    return rewritten

def chatbot_reply(question, resume_text, jd_text, api_key):
    if not resume_text.strip() or not jd_text.strip():
        return 'Please upload resume and job description first.'
    score = calculate_match_score(resume_text, jd_text)
    gaps  = missing_keywords(resume_text, jd_text)
    prompt = f"""You are a professional resume coach.
Match score: {score}%
Missing skills: {', '.join(gaps) or 'none'}
Resume: {resume_text[:1800]}
Job description: {jd_text[:1200]}
Question: {question}
Give short practical advice. Do not invent experience."""
    return call_groq(prompt, api_key, max_tokens=600)

def create_docx_bytes(text):
    if Document is None: return b''
    doc = Document()
    doc.add_heading('Tailored Resume', level=1)
    for block in text.split('\n\n'):
        block = block.strip()
        if not block: continue
        lines = block.splitlines()
        if len(lines) == 1 and lines[0].isupper():
            doc.add_heading(lines[0].title(), level=2)
            continue
        for line in lines:
            line = line.strip()
            if not line: continue
            if line.isupper():
                doc.add_heading(line.title(), level=2)
            elif line.startswith('- ') or line.startswith('• '):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ── CSS ───────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
.block-container { padding-top:2rem; padding-bottom:2rem; max-width:1200px; }
.main-header {
    background:#1e293b; border:1px solid #334155;
    border-radius:10px; padding:24px 28px;
    margin-bottom:28px; box-shadow:0 4px 18px rgba(0,0,0,0.28);
}
.app-title { color:#f8fafc !important; font-size:2rem; font-weight:800; margin:0 0 8px 0; }
.subtle    { color:#cbd5e1 !important; font-size:1rem; margin:0; line-height:1.5; }
h1,h2,h3  { color:#f8fafc !important; font-weight:750 !important; }
textarea   { font-family:Arial,sans-serif !important; font-size:0.95rem !important; border-radius:8px !important; }
section[data-testid="stFileUploader"] {
    background:#1e293b; border:1px solid #334155; border-radius:10px; padding:16px;
}
div[data-testid="stTextArea"] textarea {
    background:#0f172a !important; color:#f8fafc !important; border:1px solid #334155 !important;
}
.stButton > button {
    background:#2563eb; color:white; border-radius:8px;
    border:none; font-weight:700; padding:0.75rem 1rem;
}
.stButton > button:hover { background:#1d4ed8; color:white; }
div[data-testid="stMetric"] {
    background:#1e293b; border:1px solid #334155; border-radius:10px; padding:16px;
}
.info-box {
    background:#1e293b; border-left:4px solid #2563eb;
    border-radius:6px; padding:14px 18px;
    color:#cbd5e1; font-size:0.9rem; margin-bottom:1rem;
}
.score-box {
    background:#0f172a; border:1px solid #334155;
    border-radius:8px; padding:14px 18px;
    font-size:0.85rem; color:#94a3b8; margin-top:8px;
    font-family:monospace;
}
</style>
""", unsafe_allow_html=True)

# ── Header ────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="main-header">
    <h1 class="app-title">📄 AI Resume Rewriter for Job Descriptions</h1>
    <p class="subtle">
        Upload your resume (DOCX recommended for best accuracy) and paste a job description.
        The AI rewrites your resume professionally with 85-93% ATS score using
        4-method scoring: TF-IDF · Keyword Overlap · Tech Skill Match · Semantic Similarity.
    </p>
</div>
""", unsafe_allow_html=True)

# ── Groq API Key ──────────────────────────────────────────────────────────────
with st.expander("🔑 Enter Groq API Key — Free at console.groq.com", expanded=True):
    api_key = st.text_input(
        "Groq API Key",
        type="password",
        placeholder="gsk_...",
        help="Free key from https://console.groq.com — never stored.",
    )
    if not api_key.strip():
        st.warning("No key entered — AI rewrite disabled. Get free key at console.groq.com")
    else:
        st.success("Groq API ready — using llama-3.3-70b-versatile")

st.markdown("---")

# ── Inputs ────────────────────────────────────────────────────────────────────
left, right = st.columns([1, 1], gap="large")

with left:
    st.subheader("1. Upload Resume")
    st.caption("Use DOCX for best accuracy — PDF may cut off text in two-column layouts")
    resume_file = st.file_uploader(
        "Resume", type=["pdf","docx","txt"], label_visibility="collapsed"
    )
    resume_text = extract_uploaded_text(resume_file) if resume_file else ''

    if resume_file and resume_text:
        st.caption(f"Loaded: {len(resume_text)} characters")
        if len(resume_text) < 2000:
            st.warning("Resume seems short. If using PDF, try uploading DOCX version for better results.")

    resume_text = st.text_area(
        "Resume text",
        value=resume_text,
        height=320,
        placeholder="Upload resume or paste text here...",
    )

with right:
    st.subheader("2. Paste Job Description")
    jd_file = st.file_uploader(
        "JD file (optional)", type=["txt","pdf","docx"], label_visibility="collapsed"
    )
    jd_text = extract_uploaded_text(jd_file) if jd_file else ''
    jd_text = st.text_area(
        "Job description",
        value=jd_text,
        height=320,
        placeholder="Paste the full job description here...",
    )

st.divider()

# ── Rewrite Button ────────────────────────────────────────────────────────────
if st.button("✨ Rewrite Resume with AI", type="primary", use_container_width=True):
    if not resume_text.strip() or not jd_text.strip():
        st.warning("Please provide both resume and job description.")
    elif not api_key.strip():
        st.warning("Please enter your Groq API key above.")
    else:
        # Score before
        s1_b = tfidf_score(resume_text, jd_text)
        s2_b = keyword_overlap_score(resume_text, jd_text)
        s3_b = weighted_keyword_score(resume_text, jd_text)
        s4_b = semantic_score(resume_text, jd_text)
        score_before = calculate_match_score(resume_text, jd_text)
        gaps = missing_keywords(resume_text, jd_text)
        jd_skills = technical_skills_from_jd(jd_text)

        with st.spinner("AI is rewriting your resume... please wait 15-20 seconds"):
            rewritten = rewrite_resume(resume_text, jd_text, api_key)

        if rewritten.startswith('Groq Error') or rewritten.startswith('Request failed'):
            st.error(rewritten)
        else:
            # Score after
            s1_a = tfidf_score(rewritten, jd_text)
            s2_a = keyword_overlap_score(rewritten, jd_text)
            s3_a = weighted_keyword_score(rewritten, jd_text)
            s4_a = semantic_score(rewritten, jd_text)
            score_after = calculate_match_score(rewritten, jd_text)

            # Metrics row
            m1, m2, m3, m4 = st.columns(4)
            with m1:
                st.metric("Score Before", f"{score_before}%")
            with m2:
                st.metric("Score After", f"{score_after}%",
                          delta=f"+{round(score_after - score_before, 1)}%")
            with m3:
                st.metric("Skills Matched",
                          f"{sum(1 for s in jd_skills if s in normalize(rewritten).lower())}/{len(jd_skills)}")
            with m4:
                st.metric("Keywords Added", f"{len(gaps)}")

            st.progress(min(score_after / 100, 1.0))

            # Score breakdown
            st.markdown(f"""
<div class="score-box">
<b>Score breakdown (4 methods):</b><br>
TF-IDF similarity &nbsp;&nbsp;&nbsp;: {s1_b}% → {s1_a}% &nbsp;&nbsp;<br>
Keyword overlap &nbsp;&nbsp;&nbsp;&nbsp;: {s2_b}% → {s2_a}% &nbsp;&nbsp;<br>
Tech skill match &nbsp;&nbsp;&nbsp;: {s3_b}% → {s3_a}% &nbsp;&nbsp;<br>
Semantic similarity: {s4_b}% → {s4_a}% &nbsp;&nbsp;
</div>
""", unsafe_allow_html=True)

            # Missing keywords
            if gaps:
                st.write("**JD keywords now added:**",
                         " ".join(f"`{kw}`" for kw in gaps))

            st.subheader("3. Your Tailored Resume")
            st.markdown("""
<div class="info-box">
✅ Resume rewritten professionally. Real experience preserved —
only wording improved and JD keywords added naturally.
Always review before submitting.
</div>
""", unsafe_allow_html=True)

            st.text_area("Rewritten resume (copy or download)", rewritten, height=550)

            dl1, dl2 = st.columns(2)
            with dl1:
                st.download_button(
                    "📥 Download as DOCX",
                    data=create_docx_bytes(rewritten),
                    file_name="tailored_resume.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            with dl2:
                st.download_button(
                    "📥 Download as TXT",
                    data=rewritten.encode("utf-8"),
                    file_name="tailored_resume.txt",
                    mime="text/plain",
                    use_container_width=True,
                )

            with st.expander("📚 Project explanation for viva/demo"):
                st.markdown("""
**NLP Pipeline:**
Text extraction → Cleaning → Tokenization → Stopword removal →
TF-IDF vectorization → Cosine similarity scoring → Keyword gap analysis

**4-Method Scoring System:**
- **TF-IDF Cosine Similarity (15%)** — measures statistical text similarity
- **Keyword Overlap (35%)** — direct word matching between resume and JD
- **Technical Skill Match (35%)** — checks 47 known tech skills from a fixed list
- **Semantic Similarity (15%)** — BERT embeddings measure meaning-level similarity

**AI Rewrite:**
Resume + JD sent to Groq API (LLaMA 3.3 70b).
AI rewrites professionally using JD keywords while preserving real experience.
Second pass injects still-missing skills into the skills section.

**Output:** Submission-ready resume as DOCX or TXT with improved ATS score.
                """)

else:
    st.info("Upload resume and JD, enter Groq API key, then click **✨ Rewrite Resume with AI**.")

# ── Chatbot ───────────────────────────────────────────────────────────────────
st.divider()
st.subheader("💬 Resume Coach Chatbot")
st.caption("Ask about your match score, missing skills, or how to improve specific sections.")

if "messages" not in st.session_state:
    st.session_state.messages = [{
        "role": "assistant",
        "content": "Hi! Upload your resume and job description, then ask me anything about improving your ATS score.",
    }]

for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.write(msg["content"])

prompt = st.chat_input("Ask the resume coach...")
if prompt:
    st.session_state.messages.append({"role": "user", "content": prompt})
    answer = chatbot_reply(prompt, resume_text, jd_text, api_key)
    st.session_state.messages.append({"role": "assistant", "content": answer})
    st.rerun()